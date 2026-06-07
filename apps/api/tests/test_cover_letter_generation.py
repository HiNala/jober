from __future__ import annotations

import json
import os
from io import BytesIO

import pytest
from docx import Document

from jober_api.config import settings
from jober_api.models.enums import DocumentType, JobTargetStatus
from jober_api.models.llm_call import LlmCall
from jober_api.repositories.job_target import JobTargetRepository
from jober_api.repositories.resume_asset import ResumeAssetRepository
from jober_api.repositories.user_profile import UserProfileRepository
from jober_api.services.claims_index import build_claims_index
from jober_api.services.documents.ats_scoring import score_keyword_coverage
from jober_api.services.documents.claims_guard import parse_draft_payload, verify_draft_claims
from jober_api.services.documents.cover_letter_generator import (
    ClaimsRejectedError,
    generate_cover_letter,
)
from jober_api.services.documents.render_pdf import render_cover_letter_pdf
from jober_api.services.llm.gateway import BudgetExceededError, TemplateLlmProvider, assert_budget
from jober_api.storage.minio_client import ObjectStorage

pytestmark = pytest.mark.skipif(
    os.getenv("CI") != "true" and os.getenv("RUN_DB_TESTS") != "1",
    reason="requires Postgres",
)


def _resume_docx() -> bytes:
    document = Document()
    document.add_paragraph("Brian Permut — Glide Design")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI, TypeScript, React, RAG, Docker")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_stuffing_penalty_fires_on_high_density() -> None:
    stuffed = " ".join(["TypeScript"] * 40 + ["product engineering context"] * 5)
    report = score_keyword_coverage(
        stuffed,
        "TypeScript React Python FastAPI",
        "TypeScript React",
    )
    assert report.stuffing_penalty > 0
    assert report.density > 0.12


def test_claims_guard_rejects_injected_false_credential() -> None:
    claims = build_claims_index("Python FastAPI React", {"skills": ["Python", "React"]})
    draft = parse_draft_payload(
        json.dumps(
            {
                "body": "I am CKA certified.",
                "asserted_facts": ["Certified Kubernetes Administrator"],
                "paragraph_grounding": [],
            }
        )
    )
    result = verify_draft_claims(draft, claims)
    assert not result.ok
    assert "Certified Kubernetes Administrator" in result.unsupported


def test_pdf_renderer_produces_selectable_text_bytes() -> None:
    pdf = render_cover_letter_pdf(
        body="Hello team.\n\nI build with Python.",
        applicant_name="Brian",
        company="Acme",
        role="Engineer",
    )
    assert pdf.startswith(b"%PDF")


@pytest.mark.asyncio
async def test_generate_cover_letter_persists_document(
    db_session,
    truncate_tables,
    monkeypatch,
) -> None:
    async def _fake_put(self, key, data, content_type="application/octet-stream", length=None):
        from jober_api.storage.minio_client import StoredObject

        return StoredObject(bucket="test", key=key, etag="1")

    monkeypatch.setattr(ObjectStorage, "put_object", _fake_put)
    monkeypatch.setattr(settings, "llm_provider", "template")
    monkeypatch.setattr(settings, "llm_monthly_budget_usd", 100.0)

    profiles = UserProfileRepository(db_session)
    await profiles.create(name="Brian Permut", email="brian@example.com")
    jobs = JobTargetRepository(db_session)
    job = await jobs.create(
        company="Acme AI",
        role="Staff Engineer",
        fit_lane="AI platform",
        cover_letter_hook="Built agent stack end-to-end",
        why_fit="Strong fit",
        status=JobTargetStatus.NEW,
    )
    resumes = ResumeAssetRepository(db_session)
    text = "Brian built Python FastAPI systems with RAG and Docker at Glide."
    skills_index = {
        "skills": ["Python", "FastAPI", "RAG", "Docker"],
        "claims_index": build_claims_index(
            text,
            {"skills": ["Python", "FastAPI", "RAG", "Docker"]},
        ),
    }
    resume = await resumes.create(
        object_key="resumes/test/resume.docx",
        original_filename="resume.docx",
        extracted_text=text,
        skills_index=skills_index,
        is_active=True,
    )
    profile = await profiles.get_singleton()
    assert profile is not None
    profile.default_resume_asset_id = resume.id
    await db_session.commit()

    storage = ObjectStorage()
    result = await generate_cover_letter(
        db_session,
        storage,
        job_target_id=job.id,
        force=True,
    )
    await db_session.commit()

    assert result["document_type"] == DocumentType.COVER_LETTER.value
    assert result["ats_score"] is not None
    assert 200 <= len(result["text"].split()) <= 500
    assert result["object_key_pdf"]


@pytest.mark.asyncio
async def test_budget_blocks_generation(db_session, truncate_tables, monkeypatch) -> None:
    monkeypatch.setattr(settings, "llm_monthly_budget_usd", 0.01)
    db_session.add(
        LlmCall(
            agent_role="test",
            provider="template",
            model="test",
            cost_usd=0.02,
        )
    )
    await db_session.flush()
    with pytest.raises(BudgetExceededError):
        await assert_budget(db_session, projected_cost=0.001)


class _BadFactsProvider(TemplateLlmProvider):
    async def complete(self, **kwargs):
        completion = await super().complete(**kwargs)
        data = json.loads(completion.content)
        data["asserted_facts"] = ["Certified Kubernetes Administrator"]
        from dataclasses import replace

        return replace(completion, content=json.dumps(data))


@pytest.mark.asyncio
async def test_generation_rejects_unsupported_claims_after_retries(
    db_session,
    truncate_tables,
    monkeypatch,
) -> None:
    async def _fake_put(self, key, data, content_type="application/octet-stream", length=None):
        from jober_api.storage.minio_client import StoredObject

        return StoredObject(bucket="test", key=key, etag="1")

    monkeypatch.setattr(ObjectStorage, "put_object", _fake_put)
    monkeypatch.setattr(settings, "llm_monthly_budget_usd", 100.0)

    import jober_api.services.documents.cover_letter_generator as gen_mod

    monkeypatch.setattr(gen_mod, "get_llm_provider", lambda: _BadFactsProvider())

    jobs = JobTargetRepository(db_session)
    job = await jobs.create(company="Acme", role="Eng", status=JobTargetStatus.NEW)
    resumes = ResumeAssetRepository(db_session)
    text = "Python developer"
    await resumes.create(
        object_key="k",
        original_filename="r.docx",
        extracted_text=text,
        skills_index={
            "skills": ["Python"],
            "claims_index": build_claims_index(text, {"skills": ["Python"]}),
        },
        is_active=True,
    )
    await db_session.commit()

    with pytest.raises(ClaimsRejectedError):
        await generate_cover_letter(
            db_session,
            ObjectStorage(),
            job_target_id=job.id,
            force=True,
        )
