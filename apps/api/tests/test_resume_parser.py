from __future__ import annotations

from io import BytesIO

from docx import Document

from jober_api.services.claims_index import build_claims_index, validate_claims
from jober_api.services.resume_parser import extract_docx_text, parse_skills_index


def _sample_docx() -> bytes:
    document = Document()
    document.add_paragraph("Brian Permut")
    document.add_paragraph("Skills")
    document.add_paragraph("Python, FastAPI, PostgreSQL")
    document.add_paragraph("Experience")
    document.add_paragraph("Built agent platforms.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_extract_docx_and_parse_skills() -> None:
    data = _sample_docx()
    text = extract_docx_text(data)
    assert "Brian Permut" in text
    index = parse_skills_index(text)
    assert "Python" in index["skills"]
    assert "FastAPI" in index["skills"]


def test_claims_check_rejects_invented_credential() -> None:
    data = _sample_docx()
    text = extract_docx_text(data)
    skills = parse_skills_index(text)
    claims = build_claims_index(text, skills)
    unsupported = validate_claims(
        claims,
        ["Python", "Certified Kubernetes Administrator"],
    )
    assert "Certified Kubernetes Administrator" in unsupported
    assert "Python" not in unsupported
