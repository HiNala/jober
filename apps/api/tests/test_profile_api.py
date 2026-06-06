from __future__ import annotations

import os
from io import BytesIO

import pytest
from docx import Document
from httpx import ASGITransport, AsyncClient

from jober_api.main import app

pytestmark = pytest.mark.skipif(
    os.getenv("CI") != "true" and os.getenv("RUN_DB_TESTS") != "1",
    reason="requires Postgres",
)


def _resume_docx() -> bytes:
    document = Document()
    document.add_paragraph("Skills")
    document.add_paragraph("TypeScript, React")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_profile_vault_round_trip(
    db_session, truncate_tables, vault_key, monkeypatch
) -> None:
    from jober_api.db import session as db_session_module
    from jober_api.storage.minio_client import ObjectStorage

    async def _override():
        yield db_session

    async def _fake_put(self, key, data, content_type="application/octet-stream", length=None):
        from jober_api.storage.minio_client import StoredObject

        return StoredObject(bucket="test", key=key, etag="test")

    monkeypatch.setattr(ObjectStorage, "put_object", _fake_put)
    app.dependency_overrides[db_session_module.get_session] = _override
    try:
        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as client:
            patch = await client.patch(
                "/api/profile/vault",
                json={
                    "veteran_status": "not_a_veteran",
                    "field_consent": {
                        "veteran_status": {"consent": True, "never_autofill": False},
                    },
                },
            )
            assert patch.status_code == 200
            body = patch.json()
            veteran = next(f for f in body["fields"] if f["key"] == "veteran_status")
            assert veteran["value"] == "not_a_veteran"

            upload = await client.post(
                "/api/resumes",
                files={
                    "file": (
                        "resume.docx",
                        _resume_docx(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
            assert upload.status_code == 200
            resume = upload.json()
            assert "TypeScript" in resume["skills"]
    finally:
        app.dependency_overrides.clear()
