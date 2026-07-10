from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt

from jober_api.services.documents.letter_styles import normalize_template


def _font_size_for_template(template: str) -> int:
    key = normalize_template(template)
    if key == "compact":
        return 10
    if key == "modern":
        return 11
    return 11


def render_cover_letter_docx(
    *,
    body: str,
    applicant_name: str,
    company: str,
    role: str,
    template: str = "classic",
) -> bytes:
    document = Document()
    size = _font_size_for_template(template)
    name_para = document.add_paragraph(applicant_name)
    if name_para.runs:
        name_para.runs[0].bold = True
        name_para.runs[0].font.size = Pt(size)
    document.add_paragraph(f"Re: {role} at {company}")
    document.add_paragraph("")
    for paragraph in body.split("\n\n"):
        text = paragraph.strip()
        if text:
            para = document.add_paragraph(text)
            if para.runs:
                para.runs[0].font.size = Pt(size)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_resume_docx(
    *,
    body: str,
    applicant_name: str,
    target_role: str | None = None,
    target_company: str | None = None,
) -> bytes:
    """Resume DOCX — name header + body sections (not cover-letter 'Re:' framing)."""
    document = Document()
    name_para = document.add_paragraph(applicant_name)
    if name_para.runs:
        name_para.runs[0].bold = True
        name_para.runs[0].font.size = Pt(16)
    meta_bits = [b for b in (target_role, target_company) if b]
    if meta_bits:
        meta = document.add_paragraph(" · ".join(meta_bits))
        if meta.runs:
            meta.runs[0].font.size = Pt(9)
    document.add_paragraph("")
    for block in body.split("\n\n"):
        text = block.strip()
        if not text:
            continue
        lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
        if len(lines) == 1 and (lines[0].isupper() or (len(lines[0]) < 40 and ":" not in lines[0])):
            para = document.add_paragraph(lines[0].title() if lines[0].isupper() else lines[0])
            if para.runs:
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(11)
            continue
        para = document.add_paragraph(text)
        if para.runs:
            para.runs[0].font.size = Pt(10)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
